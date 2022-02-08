import re
import numpy as np
import pandas as pd
import openpyxl as opx
import PySimpleGUI as sg
from docx import Document
from datetime import datetime

#Sets the attributes and objects for the "Browse File" Window.

sg.theme('DarkAmber')

layout = [
    [sg.Text('Select the Daily Absenteeism Analyzer File: '), sg.FileBrowse(key='_agents_database_file_')],
    [sg.Text('Select an Agent TimeCard File: '), sg.FileBrowse(key='_agents_timecard_file_')],
    [sg.Text("Today is:"), sg.Combo(values=['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday'], key='_LISTDAYS_')],
          [sg.Button('Ok'), sg.Button('Cancel')]]

window = sg.Window('Daily Absenteeism Analyzer', layout)

data_absents = []
agents_absent = {}

#Handles the schedule file based on file format (.csv or .xlsx). This file is used as the main database

def schedule_file_handler(timecard_path, day):
    file_ext = timecard_path.split('.')[1]
    
    if file_ext == "xlsx":
        schedule_file = pd.read_excel(timecard_path, 'Schedule Checker')
        schedule_file['TO ID'] = (schedule_file['TO ID'].fillna(0)).astype(int)
        schedule_file['TO ID'] = schedule_file['TO ID'].astype(str)
        agent_schedule_dict = dict(zip(schedule_file['TO ID'], list(zip(schedule_file['IRIS ID'], schedule_file['Last Name'], schedule_file['Name'], schedule_file['Market'], schedule_file[day]))))
        return agent_schedule_dict

    elif file_ext == "csv":
        schedule_file = pd.read_csv(timecard_path, 'Schedule Checker')
        schedule_file['TO ID'] = (schedule_file['TO ID'].fillna(0)).astype(int)
        schedule_file['TO ID'] = schedule_file['TO ID'].astype(str)
        agent_schedule_dict = dict(zip(schedule_file['TO ID'], list(zip(schedule_file['IRIS ID'], schedule_file['Last Name'], schedule_file['Name'], schedule_file['Market'], schedule_file[day]))))
        return agent_schedule_dict
    else:
        sg.Popup("Please, select a .csv or .xlsx type of file!") 

#Handles the timecard file (input) based on file format (.csv or .xlsx). This file is the input file that contains all agent logins.

def timecard_file_handler(timecard_path):
    file_ext = timecard_path.split('.')[1]
    if file_ext == "xlsx":
        timecard_file = pd.read_excel(timecard_path)
        timecard_file[['Agent Name', 'ID']] = timecard_file['Agent Name (ID)'].str.split('(',expand=True)
        timecard_file['ID'] = (timecard_file['ID'].str.replace(r"\)", "")).astype(str)
        timecard_file['Time'] = timecard_file['Login Date'].dt.time
        agent_timecard_dict = dict(zip(timecard_file['ID'], list(zip(timecard_file['Agent Name'] , timecard_file['Time']))))     
        return agent_timecard_dict

    elif file_ext == "csv":
        timecard_file = pd.read_csv(timecard_path)
        timecard_file[['Agent Name', 'ID']] = timecard_file['Agent Name (ID)'].str.split('(',expand=True)
        timecard_file['ID'] = timecard_file['ID'].str.replace(r"\)", "")
        timecard_file['Time'] = timecard_file['Login Date'].dt.time
        agent_timecard_dict = dict(zip(timecard_file['ID'], list(zip(timecard_file['Agent Name'] , timecard_file['Login Date']))))   
        return agent_timecard_dict
    else:
        sg.Popup("Please, select a .csv or .xlsx type of file!")  


#Formats the absent agents dictionary (input) to be then displayed as a PySimpleGUI row in the main tab 

def agents_absent_data_formatter(agents_absent):
        
    for to_id, data in agents_absent.items():
        if(agents_absent is not None):
            row = []
            row.append(to_id)
            row.append(data[0])
            row.append(data[1])
            row.append(data[2])
            row.append(data[3])
            if(type(row) != "None"):
                data_absents.append(row)          
            else:
                pass 
        else:
            pass 

    return data_absents


#Locates the output .xlsx in the Schedule 'Checker sheet', inside the Daily Abs Analyzer file.

def update_abs_file(iris_id_absents, day, schedule_database_path):
    
    output_file = schedule_database_path

    wb = opx.load_workbook(output_file)

    ws = wb['Schedule Checker']

            
    if(day == 'Monday'):
        for row in range(2, ws.max_row+1):
            iris_id_cell = ws['B{}'.format(row)].value
            day_cell = ws['G{}'.format(row)].value
            st = ['Call Off- No repay', 'Prophylactic isolation', 'Sick leave', 'Holidays', 'Rest Day', 'Vacation - Paid']
            if(iris_id_cell in iris_id_absents and (day_cell not in st or day_cell is not None)):
                ws['G{}'.format(row)] = "Unjustified absence"
            else:
                pass
      
    elif(day == 'Tuesday'):
        for row in range(2, ws.max_row+1):
            iris_id_cell = ws['B{}'.format(row)].value
            day_cell = ws['H{}'.format(row)].value
            st = ['Call Off- No repay', 'Prophylactic isolation', 'Sick leave', 'Holidays', 'Rest Day', 'Vacation - Paid']
            if(iris_id_cell in iris_id_absents and (day_cell not in st or day_cell is not None)):
                ws['H{}'.format(row)] = "Unjustified absence"
            else:
                pass
    elif(day == 'Wednesday'):
        for row in range(2, ws.max_row+1):
            iris_id_cell = ws['B{}'.format(row)].value
            day_cell = ws['I{}'.format(row)].value
            st = ['Call Off- No repay', 'Prophylactic isolation', 'Sick leave', 'Holidays', 'Rest Day', 'Vacation - Paid']
            if(iris_id_cell in iris_id_absents and (day_cell not in st or day_cell is not None)):
                ws['I{}'.format(row)] = "Unjustified absence"
            else:
                pass
    elif(day == 'Thursday'):
        for row in range(2, ws.max_row+1):
            iris_id_cell = ws['B{}'.format(row)].value
            day_cell = ws['J{}'.format(row)].value
            st = ['Call Off- No repay', 'Prophylactic isolation', 'Sick leave', 'Holidays', 'Rest Day', 'Vacation - Paid']
            if(iris_id_cell in iris_id_absents and (day_cell not in st or day_cell is not None)):
                ws['J{}'.format(row)] = "Unjustified absence"
            else:
                pass
    elif(day == 'Friday'):
        for row in range(2, ws.max_row+1):
            iris_id_cell = ws['B{}'.format(row)].value
            day_cell = ws['K{}'.format(row)].value
            st = ['Call Off- No repay', 'Prophylactic isolation', 'Sick leave', 'Holidays', 'Rest Day', 'Vacation - Paid']
            if(iris_id_cell in iris_id_absents and (day_cell not in st or day_cell is not None)):
                ws['K{}'.format(row)] = "Unjustified absence"
            else:
                pass

    
    wb.save(output_file)       

   
    
   
#Outputs the absent agents and absenteeism main rates to a formatted word document.

def to_word_doc(rows, day):
    
    doc = Document()
    
    doc.add_heading('Daily Absenteeism of {}'.format(datetime.today().strftime('%d/%m/%Y'), 0))
    
    for info in rows:
        if info[3] == 'Dach':
            doc.add_heading('DACH', level=5)
            doc.add_paragraph(str(info[2]) + " " + str(info[1]) + "(" + str(info[0]) +") - " + str(info[4]) , style='List Bullet')            
        elif info[3] == 'Benelux':
            doc.add_heading('BENELUX', level=5)
            doc.add_paragraph(str(info[2]) + " " + str(info[1]) + "(" + str(info[0]) +") - " + str(info[4]) , style='List Bullet')            
        elif info[3] == 'France':
            doc.add_heading('FRANCE', level=5)
            doc.add_paragraph(str(info[2]) + " " + str(info[1]) + "(" + str(info[0]) +") - " + str(info[4]) , style='List Bullet')   
        elif info[3] == 'uki':
            doc.add_heading('UK&I', level=5)
            doc.add_paragraph(str(info[2]) + " " + str(info[1]) + "(" + str(info[0]) +") - " + str(info[4]) , style='List Bullet')                
        elif info[3] == 'iig':
            doc.add_heading('IIG', level=5)
            doc.add_paragraph(str(info[2]) + " " + str(info[1]) + "(" + str(info[0]) +") - " + str(info[4]) , style='List Bullet')  
        elif info[3] == 'Iberia':
            doc.add_heading('BENELUX', level=5)
            doc.add_paragraph(str(info[2]) + " " + str(info[1]) + "(" + str(info[0]) +") - " + str(info[4]) , style='List Bullet')  
        elif info[3] == 'ee':
            doc.add_heading('EE', level=5)
            doc.add_paragraph(str(info[2]) + " " + str(info[1]) + "(" + str(info[0]) +") - " + str(info[4]) , style='List Bullet') 
        elif info[3] == 'metap':
            doc.add_heading('METAP', level=5)
            doc.add_paragraph(str(info[2]) + " " + str(info[1]) + "(" + str(info[0]) +") - " + str(info[4]) , style='List Bullet') 
        else:
            doc.add_heading(str(info[3]), level=5)
            doc.add_paragraph(str(info[2]) + " " + str(info[1]) + "(" + str(info[0]) +") - " + str(info[4]) , style='List Bullet')             


    sg.theme('Dark Grey 13')

    layout_output = [[sg.Text('Choose location')],
        [sg.Input(), sg.FolderBrowse(key='_OUTPUT_WORD_DOC_')],
        [sg.OK('Run!')]
    ]

    window_output = sg.Window('Choose Location', layout_output)
    
    
    window.close()

    event, values = window_output.read()
    
    if event == 'Run!':
        doc_path = values['_OUTPUT_WORD_DOC_']+'/daily absenteeism {}.docx'.format(datetime.today().strftime('%d_%m_%Y'))
        doc.save(doc_path)
        sg.Popup("Document saved at {} !".format(doc_path))
        window_output.close()   

    


#Builds the main tab frame to display all the data collected from the rest of the functions.

def display_absents(agents_absent, day, schedule_database_path):
    
    data = agents_absent_data_formatter(agents_absent)
    sg.theme('LightGrey6')
    headings = ['TO ID', 'IRIS ID', 'Last Name', 'Name', 'Market', 'Notes']
    

    column_to_id =[ [sg.Text(d[0])] for d in data ]
    column_iris_id =[ [sg.Text(d[1])] for d in data ]
    column_last_name =[ [sg.Text(d[2])] for d in data ]
    column_name =[ [sg.Text(d[3])] for d in data ]
    column_market =[ [sg.Text(d[4])] for d in data ]
    column_notes =[ [sg.Input(key='notes_{}'.format(x))] for x in range(len(data))]
    all_columns_frame = [ [sg.Column(column_to_id),
                sg.Column(column_iris_id),
                sg.Column(column_last_name),
                sg.Column(column_name),
                sg.Column(column_market),
                sg.Column(column_notes),  
                ]        
        ]

    layout_absents = [
        [
            [sg.Text(x) for x in headings],
            [sg.Column(all_columns_frame, scrollable=True, size=(880,550))],
            [sg.Button('Export')]
        ]

    ]        
    
    window_absents = sg.Window('Daily Absenteeism Analyzer', layout_absents, size=(880,650), text_justification='left', resizable=False, finalize=True)

    while True:
        event, values = window_absents.read()
        if event == sg.WIN_CLOSED or event == "Cancel":
            break
        
        if event == "Export":
            rows = []
            headers = ['IRIS ID', 'Last Name', 'Name', 'Market', 'Notes']
            x = 0
            for info in agents_absent.values():
                rows.append([info[0], info[1], info[2], info[3], values['notes_{}'.format(x)]])
                x += 1
            df = pd.DataFrame(rows, columns=headers)
            iris_id_absents = df['IRIS ID'].tolist()
            update_abs_file(iris_id_absents, day ,schedule_database_path)
            to_word_doc(rows, day)
            window_absents.close()
        
    

#Main function
#Checks what agents from the database schedule (dict) are not in the timecard dictionary. This is to check what agents didn't login at the time of the script running

def schedule_checker(timecard_path, day, schedule_database_path):
    agent_schedule_dict = schedule_file_handler(schedule_database_path, day)   
    agent_timecard_dict = timecard_file_handler(timecard_path)
    st = ['Call Off- No repay', 'Prophylactic isolation', 'Sick leave', 'Holidays', 'Rest Day', 'Vacation - Paid', None]
    for to_id, full_name in agent_schedule_dict.items():
        try:
            if(to_id not in agent_timecard_dict) and (str(full_name[4]) not in st):
                agents_absent[to_id] = full_name
            else:
                pass
        except:
            pass
    display_absents(agents_absent, day, schedule_database_path)


while True:
    event, values = window.read()
    if event == sg.WIN_CLOSED or event == 'Cancel':
        break
    
    timecard_path = values['_agents_timecard_file_']
    schedule_database_path = values['_agents_database_file_']
    if values['_LISTDAYS_'] == "":
        sg.Popup("Please, select a day!")
    else:
        day = values['_LISTDAYS_']
        schedule_checker(timecard_path, day, schedule_database_path)


window.close()   
